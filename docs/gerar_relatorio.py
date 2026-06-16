from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

out_charts = r'C:\Users\marip\OneDrive\Área de Trabalho\fii-monitor\docs\charts'
out_file = r'C:\Users\marip\OneDrive\Área de Trabalho\fii-monitor\docs\DEVA11_Analise_Marco2026.docx'

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1a, 0x52, 0x7b)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    return p

def add_body(doc, text, bold=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)
        run.italic = True

def shade_row(row, hex_color):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

def add_alert(doc, text, tipo='alerta'):
    p = doc.add_paragraph()
    prefix = 'ALERTA: ' if tipo == 'alerta' else 'POSITIVO: '
    run = p.add_run(prefix + text)
    run.bold = True
    run.font.size = Pt(10)
    if tipo == 'alerta':
        run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
    else:
        run.font.color.rgb = RGBColor(0x27, 0xae, 0x60)

# CAPA
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ANALISE DE FII - RELATORIO GERENCIAL')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DEVA11 - Devant Recebiveis Imobiliarios FII')
run.font.size = Pt(22)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x52, 0x7b)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Relatorio Gerencial | Competencia: Marco 2026')
run.font.size = Pt(12)

doc.add_paragraph()

t = doc.add_table(rows=2, cols=5)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.style = 'Table Grid'
headers = ['Cota Mercado', 'Cota Patrimonial', 'P/VP', 'DY Mensal', 'Veredicto']
vals = ['R$ 23,90', 'R$ 94,63', '0,25x', '1,26% (mar/26)', 'MANUTENCAO COM CAUTELA']
shade_row(t.rows[0], '1a527b')
for i, (h, v) in enumerate(zip(headers, vals)):
    t.rows[0].cells[i].text = h
    t.rows[1].cells[i].text = v
    for run in t.rows[0].cells[i].paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        run.bold = True
        run.font.size = Pt(9)
    for run in t.rows[1].cells[i].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)
        if i == 4:
            run.font.color.rgb = RGBColor(0xe6, 0x7e, 0x22)

doc.add_paragraph()
add_heading(doc, 'SUMARIO EXECUTIVO', 2)
add_body(doc,
    'O DEVA11 apresenta em marco/2026 estabilizacao fragil: distribuicao mantida em R$ 0,30/cota pelo '
    'terceiro mes consecutivo, sustentada em parte pelo consumo da reserva de lucros (-R$ 349 mil). '
    'A receita gerou R$ 4,96M para distribuir R$ 4,21M - margem existe mas nao comporta deterioracao adicional.')

doc.add_paragraph()
add_alert(doc, '65,2% da carteira de CRIs em carencia de juros - mais de 2/3 do portfolio nao gera caixa agora. '
    'Este e o fator de risco mais relevante do fundo.', 'alerta')
add_alert(doc, 'Mercado precifica cota 74,7% abaixo do patrimonial (P/VP 0,25x). '
    'Mercado nao acredita que o PL e recuperavel integralmente.', 'alerta')
add_alert(doc, '9,7% da carteira inadimplente - proximo ao limiar critico para FIIs de papel.', 'alerta')
add_alert(doc, '99% da carteira protegida contra deflacao. IPCA + 10,66% ao ano - remuneracao contratada excelente se devedores pagarem.', 'positivo')

doc.add_page_break()

# PAGINA 2 - DRE
add_heading(doc, '1. ANALISE DE RESULTADOS - MARCO 2026', 1)
add_heading(doc, '1.1 DRE Gerencial - Evolucao Mensal', 2)

t2 = doc.add_table(rows=9, cols=8)
t2.style = 'Table Grid'
shade_row(t2.rows[0], '1a527b')
cols_dre = ['Item', 'set/25', 'out/25', 'nov/25', 'dez/25', 'jan/26', 'fev/26', 'mar/26']
rows_dre = [
    ('Receitas Totais (R$)', '6.405.510', '5.056.121', '6.334.352', '5.182.939', '4.879.891', '4.830.216', '4.962.852'),
    ('  Juros CRI', '5.564.241', '4.278.231', '5.167.890', '4.350.830', '3.868.740', '3.939.540', '4.063.797'),
    ('  Correc. Monet.', '275.536', '258.433', '428.361', '197.654', '364.792', '309.632', '322.851'),
    ('  Receita FIIs', '478.429', '478.429', '478.429', '478.429', '509.997', '478.429', '478.429'),
    ('  Rend. Caixa', '47.933', '41.029', '58.441', '112.080', '136.363', '102.615', '97.775'),
    ('Despesas (R$)', '(454.944)', '(434.496)', '(429.509)', '(340.250)', '(348.463)', '(402.092)', '(399.458)'),
    ('Reserva Lucros', '(1.034.848)', '13.195', '274.917', '775.274', '(266.953)', '(214.652)', '(349.921)'),
    ('Distrib./cota (R$)', 'R$0,35', 'R$0,33', 'R$0,44', 'R$0,40', 'R$0,30', 'R$0,30', 'R$0,30'),
]
for i, col in enumerate(cols_dre):
    t2.rows[0].cells[i].text = col
    for run in t2.rows[0].cells[i].paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        run.bold = True
        run.font.size = Pt(8)
for ri, row_data in enumerate(rows_dre):
    row = t2.rows[ri + 1]
    for ci, val in enumerate(row_data):
        row.cells[ci].text = val
        for run in row.cells[ci].paragraphs[0].runs:
            run.font.size = Pt(8)
            if ci == 0:
                run.bold = True

doc.add_paragraph()

chart_path = os.path.join(out_charts, 'fig1_receita_distribuicao.png')
if os.path.exists(chart_path):
    doc.add_picture(chart_path, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, 'Figura 1 - Receita total e distribuicao por cota (set/25-mar/26). Fonte: Relatorio Gerencial DEVA11, marco/2026.')

doc.add_page_break()

# PAGINA 3 - RISCOS
add_heading(doc, '2. ANALISE DE RISCOS', 1)

add_heading(doc, '2.1 Risco Central: 65,2% da Carteira em Carencia de Juros', 2)
add_body(doc,
    'Apenas 25,1% dos ativos pagam juros regularmente. A maior parte (65,2%) esta em carencia - devedores '
    'nao pagam juros agora, acumulam para pagar no futuro. Implicacoes:\n\n'
    '(a) Caixa atual sustentado por fracao da carteira - nao pelo potencial pleno do portfolio.\n\n'
    '(b) Quando carencias vencerem, pagamentos devem normalizar - mas SOMENTE se devedores tiverem '
    'capacidade financeira para honrar os juros acumulados. Com CDI a 14,75%, esse e o grande interrogante.')

chart_path = os.path.join(out_charts, 'fig3_status_cri.png')
if os.path.exists(chart_path):
    doc.add_picture(chart_path, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, 'Figura 2 - Status da carteira de CRIs. Fonte: Relatorio Gerencial DEVA11, marco/2026.')

add_heading(doc, '2.2 Concentracao em Segmentos de Alto Risco', 2)
add_body(doc,
    'Multipropriedade (33,2%) + Loteamento (22,4%) = 55,6% da carteira.\n\n'
    'Multipropriedade: venda direta a pessoa fisica com alto endividamento. Inadimplencia estrutural alta, '
    'sensivel a desemprego e credito caro. CDI a 14,75% pressiona diretamente esse perfil de comprador.\n\n'
    'Loteamento: ciclo longo (3-7 anos do lancamento ao repasse bancario). Alguns CRIs individuais mostram '
    'razao PMT executada acima do limite contratado - sinal de estresse nos pagamentos.')

chart_path = os.path.join(out_charts, 'fig4_segmentos.png')
if os.path.exists(chart_path):
    doc.add_picture(chart_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, 'Figura 3 - Composicao por segmento. Fonte: Relatorio Gerencial DEVA11, marco/2026.')

doc.add_page_break()

# PAGINA 4 - RESERVA E DESCONTO
add_heading(doc, '2.3 Reserva de Lucros: Distribuicao Sustentavel?', 2)
add_body(doc,
    'R$ 0,30/cota desde janeiro/2026 com margem estreita. Em marco: receita R$ 4,96M, distribuicao R$ 4,21M. '
    'Nos ultimos 3 meses a reserva consumiu -R$ 930 mil acumulados. Ha pressao persistente de despesas e '
    'inadimplencia que consome o excedente gerado. A distribuicao nao e artificial, mas a margem de seguranca e pequena.')

chart_path = os.path.join(out_charts, 'fig2_reserva_lucros.png')
if os.path.exists(chart_path):
    doc.add_picture(chart_path, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, 'Figura 4 - Variacao da reserva de lucros. Negativo = consumo de reserva. Fonte: Relatorio Gerencial DEVA11, marco/2026.')

add_heading(doc, '2.4 O Desconto de 74,7%: Oportunidade ou Aviso?', 2)
add_body(doc,
    'P/VP de 0,25x - um dos maiores descontos do segmento de FIIs de papel.\n\n'
    'Argumento otimista: o PL de R$ 1,33B e real - CRIs com contratos formais, garantias imobiliarias, '
    'IPCA + 10,66%. Se devedores pagarem, o fundo vale muito mais do que o mercado precifica.\n\n'
    'Argumento do mercado (pessimista): com 74,9% sem gerar caixa agora, o mercado nao acredita que o valor '
    'integral dos CRIs sera recuperado. Um CRI inadimplente em execucao de garantia vale fracao do valor de face - '
    'e 9,7% ja esta nessa situacao.\n\n'
    'Veredicto: o desconto e simultaneamente um aviso e uma precificacao racional de risco. Nao e oportunidade obvia.')

chart_path = os.path.join(out_charts, 'fig6_desconto_pvp.png')
if os.path.exists(chart_path):
    doc.add_picture(chart_path, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, 'Figura 5 - Desconto P/VP. Fonte: Relatorio Gerencial DEVA11, marco/2026.')

doc.add_page_break()

# PAGINA 5 - DY E CONCLUSAO
add_heading(doc, '3. DIVIDEND YIELD - LEITURA CRITICA', 1)
add_body(doc,
    'DY de 1,26%/mes (18,06% em 12 meses) = 103,74% CDI bruto e 122,04% gross-up IR. '
    'Matematicamente atrativo. Mas o DY elevado e resultado direto do desconto de mercado, nao de geracao '
    'de caixa operacional superior. Quem comprou a R$ 50 recebe 0,72%/mes - abaixo do CDI. '
    'O DY de 1,26% so vale para quem comprar hoje a R$ 23,90.\n\n'
    'A pergunta correta nao e "o DY e alto?" mas: "os R$ 0,30/cota sao sustentaveis?"')

chart_path = os.path.join(out_charts, 'fig5_dy_mensal.png')
if os.path.exists(chart_path):
    doc.add_picture(chart_path, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, 'Figura 6 - DY Mensal DEVA11. Fonte: Relatorio Gerencial DEVA11, marco/2026.')

add_heading(doc, '3.1 Cenarios de Sustentabilidade', 2)

t3 = doc.add_table(rows=4, cols=3)
t3.style = 'Table Grid'
shade_row(t3.rows[0], '1a527b')
cenarios = [
    ('Cenario', 'Condicao', 'Distribuicao Estimada'),
    ('Base (atual)', '65% em carencia + 9,7% inadimplente', 'R$ 0,28-0,32/cota'),
    ('Positivo', 'Carencias vencem, devedores normalizam', 'R$ 0,38-0,45/cota'),
    ('Negativo', 'Inadimplencia sobe para 15%+', 'R$ 0,20-0,25/cota ou corte'),
]
for ri, row_data in enumerate(cenarios):
    for ci, val in enumerate(row_data):
        t3.rows[ri].cells[ci].text = val
        for run in t3.rows[ri].cells[ci].paragraphs[0].runs:
            run.font.size = Pt(9)
            if ri == 0:
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                run.bold = True

doc.add_paragraph()

add_heading(doc, '4. AVALIACAO FINAL - VALE MANTER?', 1)

t4 = doc.add_table(rows=7, cols=3)
t4.style = 'Table Grid'
shade_row(t4.rows[0], '2c3e50')
fatores = [
    ('Fator', 'Leitura', 'Impacto'),
    ('Remuneracao contratada', 'IPCA + 10,66% ao ano', 'POSITIVO'),
    ('65,2% em carencia', 'Sem caixa de 2/3 da carteira agora', 'RISCO ALTO'),
    ('9,7% inadimplencia', 'Proximo ao limiar critico', 'RISCO MODERADO-ALTO'),
    ('Desconto 74,7%', 'Mercado precifica recuperacao parcial do PL', 'ALERTA ESTRUTURAL'),
    ('Reserva consumida', 'Distribuicao com margem estreita', 'RISCO MODERADO'),
    ('Multipropriedade 33%', 'Segmento mais fragil do portfolio', 'RISCO CONCENTRADO'),
]
for ri, row_data in enumerate(fatores):
    for ci, val in enumerate(row_data):
        t4.rows[ri].cells[ci].text = val
        for run in t4.rows[ri].cells[ci].paragraphs[0].runs:
            run.font.size = Pt(9)
            if ri == 0:
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                run.bold = True
            if 'POSITIVO' in val:
                run.font.color.rgb = RGBColor(0x27, 0xae, 0x60)
                run.bold = True
            elif 'RISCO' in val or 'ALERTA' in val:
                run.font.color.rgb = RGBColor(0xe6, 0x7e, 0x22)
                run.bold = True

doc.add_paragraph()
add_heading(doc, 'Conclusao', 2)
add_body(doc,
    'DEVA11 e um FII de papel em estagio de estresse moderado-alto, nao em recuperacao clara. '
    'A distribuicao de R$ 0,30/cota e sustentavel no curto prazo, mas depende de:\n\n'
    '(1) CRIs em carencia normalizarem pagamentos ao vencer os periodos de carencia;\n'
    '(2) Inadimplencia nao subir dos atuais 9,7% para 15%+ - nivel que comprometeria a distribuicao;\n'
    '(3) Macroeconomia nao piorar (juros altos prolongados aumentam inadimplencia em multipropriedade e loteamento).\n\n'
    'Para o investidor que tem DEVA11 comprado a precos mais altos: manter exposicao a um FII com '
    'esse nivel de desconto e esses riscos estruturais so faz sentido com convicção de recuperacao. '
    'Nao ha sinal claro de virada no relatorio de marco. A gestora monitora os ativos sem divulgar '
    'cronograma de normalizacao das carencias ou plano concreto de reducao de inadimplencia.\n\n'
    'VEREDICTO: MANUTENCAO COM CAUTELA.\n'
    'Observe os proximos 2 relatorios com foco em: (a) mudanca no % em carencia, '
    '(b) evolucao da inadimplencia, (c) variacao da reserva de lucros.')

doc.add_paragraph()
p_disc = doc.add_paragraph(
    'DISCLAIMER: Este relatorio e baseado exclusivamente nos documentos publicos do DEVA11 '
    '(Informe Mensal CVM Abril/2026 e Relatorio Gerencial Marco/2026). '
    'Nao constitui recomendacao de investimento. Dados de mercado referentes a 12/06/2026 conforme extrato B3.')
for run in p_disc.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)
    run.italic = True

doc.save(out_file)
print('Relatorio salvo:', out_file)
